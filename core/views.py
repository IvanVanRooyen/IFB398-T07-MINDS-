# core/views.py
from __future__ import annotations
from pydoc import doc

from django.apps import apps
from django.contrib.auth.decorators import login_required
from django.views.decorators.http import require_POST
from django.core.paginator import Paginator
from django.db.models import Q
from django.core.exceptions import PermissionDenied
from django.http import Http404, HttpResponse, JsonResponse, HttpResponseBadRequest
from django.urls import reverse
from django.views.decorators.http import require_GET, require_http_methods
from django.contrib import messages
from django.core.cache import cache
from django.shortcuts import render, get_object_or_404, redirect
from core.ai.report_service import generate_project_report
from .ai.granite_client import GraniteClient

from types import SimpleNamespace
import logging

# Exporting report
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.enums import TA_LEFT
from docx import Document as DocxDocument
from docx.shared import Pt, RGBColor
import re, io

from django.contrib.contenttypes.models import ContentType
from .forms import DocumentForm, DocumentSearchForm
from .models import Document, Process, SavedReport, AuditLog, log_audit, Prospect, DocLink, UserProfile
from .permissions import role_required, clearance_required, log_view_access
from .utils import sha256_file, extract_text, chunk_text

from .tagging import TAG_LABEL


# ---------- Helpers ----------


def _get_model(app_label: str, model_name: str):
    """
    Best-effort dynamic model fetch (lets views work even if model doesn’t exist yet).
    """
    try:
        return apps.get_model(app_label, model_name)
    except LookupError:
        return None


def _count_model(app_label: str, model_name: str, where_clause: str = None) -> int:
    mdl = _get_model(app_label, model_name)
    if mdl is None:
        return 0
    return mdl.objects.count()


def _paginate(queryset, request, per_page: int = 20):
    paginator = Paginator(queryset, per_page)
    page_number = request.GET.get("page")
    return paginator.get_page(page_number)


def _org_qs_filter(request):
    """
    Returns a Q object for organisation-scoped queryset filtering.
    - Superusers: Q() — no restriction, see all data.
    - Authenticated users with an assigned organisation: Q(organisation=their_org).
    - Authenticated users with no organisation: Q(pk__in=[]) — see nothing.
    """
    if request.user.is_superuser:
        return Q()
    if request.user.is_authenticated and hasattr(request.user, 'profile'):
        org = request.user.profile.organisation
        if org is not None:
            return Q(organisation=org)
    return Q(pk__in=[])


# ---------- Landing / Dashboard ----------


@login_required
@require_GET
def home(request):
    """
    Simple landing that shows recent Projects & Documents (as per your snippet).
    """
    org_filter = _org_qs_filter(request)
    projects = Process.objects.filter(org_filter).order_by("-created_at")[:10]
    docs = Document.objects.filter(org_filter).order_by("-created_at")[:10]
    return render(
        request,
        "core/home.html",
        {"projects": projects, "docs": docs},
    )


@login_required
@require_GET
def dashboard(request):
    """
    Dashboard cards + quick links. Works even if domain models aren’t ready yet.
    """
    org_filter = _org_qs_filter(request)
    Prospect = _get_model("core", "Prospect")
    Drillhole = _get_model("core", "Drillhole")
    Tenement = _get_model("core", "Tenement")
    metrics = {
        "project_count": Process.objects.filter(org_filter).count(),
        "document_count": Document.objects.filter(org_filter).count(),
        "prospect_count": Prospect.objects.filter(org_filter).count() if Prospect else 0,
        "drillhole_count": Drillhole.objects.filter(org_filter).count() if Drillhole else 0,
        "tenement_count": Tenement.objects.filter(org_filter).count() if Tenement else 0,
    }
    recent_docs = Document.objects.filter(org_filter).order_by("-created_at")[:8]
    return render(
        request,
        "core/dashboard.html",
        {"metrics": metrics, "recent_docs": recent_docs},
    )


# Optional: HTMX endpoint to refresh stats without reloading the whole page
@login_required
@require_GET
def stats_partial(request):
    org_filter = _org_qs_filter(request)
    Prospect = _get_model("core", "Prospect")
    Drillhole = _get_model("core", "Drillhole")
    Tenement = _get_model("core", "Tenement")
    ctx = {
        "project_count": Process.objects.filter(org_filter).count(),
        "document_count": Document.objects.filter(org_filter).count(),
        "prospect_count": Prospect.objects.filter(org_filter).count() if Prospect else 0,
        "drillhole_count": Drillhole.objects.filter(org_filter).count() if Drillhole else 0,
        "tenement_count": Tenement.objects.filter(org_filter).count() if Tenement else 0,
    }
    return render(request, "core/partials/stats.html", ctx)


# ---------- Cache keys ----------

DOCS_CACHE_KEY = "docs:unfiltered:page1:v1"
DOCS_CACHE_TTL = 120  # 2 minutes


def _docs_cache_key(request):
    """Per-organisation cache key so users only see their own org's cached documents."""
    if request.user.is_superuser:
        return "docs:unfiltered:page1:v1:all"
    org = None
    if request.user.is_authenticated and hasattr(request.user, 'profile'):
        org = request.user.profile.organisation
    org_id = str(org.id) if org else "noorg"
    return f"docs:unfiltered:page1:v1:{org_id}"


# ---------- Documents ----------

log = logging.getLogger(__name__)


def _get_clearance_level(request) -> str:
    """return the requesting user's clearance level string, defaulting to PUBLIC"""
    if request.user.is_authenticated and hasattr(request.user, "profile"):
        return request.user.profile.clearance_level
    return "PUBLIC"


def _report_cache_key(process_id: str, clearance_level: str, latest_doc_ts) -> str:
    doc_fingerprint = latest_doc_ts.strftime("%Y%m%d%H%M%S%f") if latest_doc_ts else "empty"
    return f"report:v1:{process_id}:{clearance_level}:{doc_fingerprint}"


def _get_cached_report_md(process_id: str, clearance_level: str) -> str:
    """
    return the cached report markdown for this project + clearance combination, generating and caching if not already existing

    cache key includes a content fingerprint (latest document upload timestamp)
    so the cache self-invalidates whenever a new document is added to the project
    """
    latest_doc_ts = (
        Document.objects
        .filter(process_id=process_id)
        .order_by("-created_at")
        .values_list("created_at", flat=True)
        .first()
    )
    cache_key = _report_cache_key(process_id, clearance_level, latest_doc_ts)

    md = cache.get(cache_key)
    if md is None:
        md = generate_project_report(process_id, clearance_level=clearance_level)
        cache.set(cache_key, md, 86400)  # 24 hours
    return md

@login_required
@role_required(
    UserProfile.RoleChoices.GEOLOGIST_EXPL,
    UserProfile.RoleChoices.FIELD_LEAD,
    UserProfile.RoleChoices.DATA_MANAGER,
    UserProfile.RoleChoices.GEOLOGIST_MINE,
    UserProfile.RoleChoices.METALLURGIST,
    UserProfile.RoleChoices.OPERATIONS_MANAGER,
    UserProfile.RoleChoices.ADMIN,
)
@require_http_methods(["GET", "POST"])
def upload_doc(request):
    """
    Upload with SHA-256 de-duplication (your original logic, with tiny polish).
    """
    if request.method == "POST":
        form = DocumentForm(request.POST, request.FILES)
        log.debug("FILES keys: %s", list(request.FILES.keys()))  # debug: ensure 'file' is present
        if form.is_valid():
            doc = form.save(commit=False)
            # Only set if user is authenticated (created_by is nullable)
            if request.user.is_authenticated:
                doc.created_by = request.user
            
            # Give extracted text a safe default in case extraction fails, to avoid null issues in search
            doc.extracted_text = ""

            if doc.file:
                # Important: call sha256_file on the uploaded file *before* saving
                doc.checksum_sha256 = sha256_file(doc.file)
                doc.extracted_text = extract_text(doc.file) or ""

            if doc.checksum_sha256 and Document.objects.filter(
                checksum_sha256=doc.checksum_sha256
            ).exists():
                # Duplicate detected — re-render with error + keep their form state
                docs = Document.objects.filter(_org_qs_filter(request)).order_by("-created_at")[:20]
                return render(
                    request,
                    "core/upload.html",
                    {
                        "form": form,
                        "docs": docs,
                        "error": "Duplicate file detected (checksum match).",
                    },
                )

            doc.extracted_text = doc.extracted_text or ""

            #Debug
            print("BEFORE SAVE extracted_text:", repr(doc.extracted_text))
            print("BEFORE SAVE type:", type(doc.extracted_text))
            print("BEFORE SAVE dict:", {
                "title": doc.title,
                "doc_type": doc.doc_type,
                "confidentiality": doc.confidentiality,
                "organisation_id": doc.organisation_id,
                "process_id": doc.process_id,
                "created_by_id": doc.created_by_id,
                "extracted_text": repr(doc.extracted_text),
            })

            doc.save()
            # form.save_m2m()
            # Build text chunks for RAG retrieval
            if doc.extracted_text:
                from .models import DocumentChunk
                chunks = chunk_text(doc.extracted_text)
                DocumentChunk.objects.bulk_create([
                    DocumentChunk(
                    document=doc,
                    chunk_index=i,
                    text=chunk,
                    process=doc.process,
                    doc_type=doc.doc_type,
                    timestamp=doc.timestamp,
                    )
                    for i, chunk in enumerate(chunks)
                ])

            # Invalidate the unfiltered document list cache so the new doc appears immediately
            cache.delete(_docs_cache_key(request))

            # Pre-warm the report cache for this project (shifts LLM wait to upload time)
            if doc.process:
                uploader_clearance = _get_clearance_level(request)
                warm_cache_key = _report_cache_key(
                    str(doc.process_id), uploader_clearance, doc.created_at
                )
                try:
                    warmed = generate_project_report(
                        str(doc.process_id), clearance_level=uploader_clearance
                    )
                    cache.set(warm_cache_key, warmed, 86400)
                except Exception:
                    # Granite unavailable — report will be generated on first view request
                    pass

            return redirect("upload")
        else:
            # Show validation errors + keep the recent docs list
            # Show *why* it failed
            log.warning("Upload invalid: %s", form.errors)
            docs = Document.objects.filter(_org_qs_filter(request)).order_by("-created_at")[:20]
            return render(
                request,
                "core/upload.html",
                {"form": form, "docs": docs, "error": "Please correct the errors below."},
            )

    # GET
    form = DocumentForm()
    docs = Document.objects.filter(_org_qs_filter(request)).order_by("-created_at")[:20]
    return render(request, "core/upload.html", {"form": form, "docs": docs})


@login_required
@require_GET
def documents(request):
    """
    Document library with search + tag + pagination.
    """
    # Build doc_type choices from whatever is actually in the DB
    existing_types = (
        Document.objects
        .filter(_org_qs_filter(request))
        .exclude(doc_type="")
        .exclude(doc_type__isnull=True)
        .values_list("doc_type", flat=True)
        .distinct()
        .order_by("doc_type")
    )

    type_choices = [("", "All types")] + [(t, t) for t in existing_types]

    form = DocumentSearchForm(request.GET or None, doc_type_choices=type_choices)
    qs = Document.objects.filter(_org_qs_filter(request)).select_related("process", "organisation").order_by("-created_at")

    q_value = ""
    if form.is_valid():

        # Full-text : title, doc_type, confidentiality, project name, org
        q = form.cleaned_data.get("q", "").strip()

        if q:
            qs = qs.filter(
                Q(title__icontains=q)
                | Q(doc_type__icontains=q)
                | Q(confidentiality__icontains=q)
                | Q(process__name__icontains=q)
                | Q(organisation__name__icontains=q)
                | Q(extracted_text__icontains=q)
            )

        # Project
        process = form.cleaned_data.get("process")
        if process:
            qs = qs.filter(process=process)
 
        # Date range (inclusive, on the document's own date)
        date_from = form.cleaned_data.get("date_from")
        if date_from:
            qs = qs.filter(timestamp__gte=date_from)
 
        date_to = form.cleaned_data.get("date_to")
        if date_to:
            qs = qs.filter(timestamp__lte=date_to)
 
        # Metadata
        doc_type = form.cleaned_data.get("doc_type")
        if doc_type:
            qs = qs.filter(doc_type__iexact=doc_type)
 
        confidentiality = form.cleaned_data.get("confidentiality")
        if confidentiality:
            qs = qs.filter(confidentiality__iexact=confidentiality)
 
        tag = form.cleaned_data.get("tag")
        if tag:
            try:
                qs = qs.filter(tags__contains=[int(tag)])
            except (TypeError, ValueError):
                pass
 
    filters_active = any(request.GET.get(f) for f in
                         ["q", "process", "date_from", "date_to", "doc_type", "confidentiality", "tag"])

    page_num = request.GET.get("page", "1")

    # Serve from cache for the default view (no filters, page 1)
    if not filters_active and page_num == "1":
        cached = cache.get(_docs_cache_key(request))
        if cached is not None:
            # Reconstruct a Page-like proxy from the cached dict so the template
            # interface (page.object_list, page.has_next, etc.) works unchanged.
            page_proxy = SimpleNamespace(
                object_list=cached["docs"],
                has_other_pages=cached["has_next"] or cached["has_previous"],
                has_previous=cached["has_previous"],
                has_next=cached["has_next"],
                number=1,
                paginator=SimpleNamespace(num_pages=cached["num_pages"]),
                previous_page_number=cached["prev_page_number"],
                next_page_number=cached["next_page_number"],
            )
            return render(request, "core/documents.html", {
                "form": form,
                "page": page_proxy,
                "q": "",
                "filters_active": False,
            })

    page = _paginate(qs, request, per_page=24)

    # cache only the unfiltered page-1 result abd Store a plain dict rather than the Page object to avoid serialising the full queryset into Redis
    if not filters_active and page_num == "1":
        cache.set(_docs_cache_key(request), {
            "docs": list(page.object_list),
            "num_pages": page.paginator.num_pages,
            "has_next": page.has_next(),
            "has_previous": page.has_previous(),
            "next_page_number": page.next_page_number() if page.has_next() else None,
            "prev_page_number": page.previous_page_number() if page.has_previous() else None,
        }, DOCS_CACHE_TTL)

    return render(request, "core/documents.html", {
        "form": form,
        "page": page,
        "q": q_value,
        "filters_active": filters_active,
    })


@login_required
@log_view_access(Document)
@require_GET
def document_detail(request, pk):
    doc = get_object_or_404(Document, pk=pk)
    if not request.user.is_superuser:
        if (
            hasattr(request.user, 'profile')
            and request.user.profile.organisation
            and doc.organisation
            and doc.organisation != request.user.profile.organisation
        ):
            raise PermissionDenied
    tag_labels = [TAG_LABEL.get(t, f"Tag {t}") for t in (doc.tags or [])]
    return render(request, "core/document_detail.html", {
        "doc": doc,
        "tag_labels": tag_labels,
    })


@login_required
@role_required(
    UserProfile.RoleChoices.FIELD_LEAD,
    UserProfile.RoleChoices.DATA_MANAGER,
    UserProfile.RoleChoices.OPERATIONS_MANAGER,
    UserProfile.RoleChoices.ADMIN,
)
@require_http_methods(["POST", "DELETE"])
def delete_document(request, pk):
    """
    Delete a document and its associated file from storage (MinIO).
    """
    doc = get_object_or_404(Document, pk=pk)
    if not request.user.is_superuser:
        if (
            hasattr(request.user, 'profile')
            and request.user.profile.organisation
            and doc.organisation
            and doc.organisation != request.user.profile.organisation
        ):
            raise PermissionDenied

    # Store title for success message
    doc_title = doc.title

    try:
        # The delete() method on the model will handle file deletion from MinIO
        doc.delete()

        # Invalidate the document list cache so the deletion is reflected immediately
        cache.delete(_docs_cache_key(request))

        # Return JSON response for HTMX/AJAX requests
        if request.headers.get('HX-Request'):
            return JsonResponse({
                "success": True,
                "message": f"Document '{doc_title}' deleted successfully."
            })

        # Redirect for regular form submissions
        return redirect("upload")

    except Exception as e:
        if request.headers.get('HX-Request'):
            return JsonResponse({
                "success": False,
                "message": f"Error deleting document: {str(e)}"
            }, status=500)

        # For regular requests, redirect with error (would need messages framework)
        return redirect("upload")


# ---------- Projects / Domain pages (safe even if models are missing) ----------


@login_required
@require_GET
def prospects(request):
    Prospect = _get_model("core", "Prospect")
    if Prospect:
        qs = Prospect.objects.filter(_org_qs_filter(request)).order_by("-created_at")
        page = _paginate(qs, request)
    else:
        qs, page = [], None
    return render(
        request,
        "core/prospects.html",
        {"page": page, "model_exists": Prospect is not None},
    )


@login_required
def prospect_detail(request, pk):
    prospect = get_object_or_404(Prospect, pk=pk)
    if not request.user.is_superuser:
        if (
            hasattr(request.user, 'profile')
            and request.user.profile.organisation
            and prospect.organisation
            and prospect.organisation != request.user.profile.organisation
        ):
            raise PermissionDenied
    doc_links = DocLink.objects.filter(
        content_type=ContentType.objects.get_for_model(Prospect),
        object_id=prospect.pk,
    ).select_related("document", "created_by").order_by("-created_at")
    return render(request, "core/prospect_detail.html", {
        "prospect": prospect,
        "doc_links": doc_links,
    })


# ---------- DocLink Views ----------

_LINKABLE_MODELS = {
    "prospect": ("core", "prospect"),
    "tenement": ("core", "tenement"),
    "drillhole": ("core", "drillhole"),
    "process": ("core", "process"),
}


@login_required
@require_http_methods(["GET"])
def doc_link_picker(request):
    """HTMX partial: render the document picker modal for linking a document to an entity."""
    content_type_label = request.GET.get("content_type", "")
    object_id = request.GET.get("object_id", "")

    if content_type_label not in _LINKABLE_MODELS:
        return HttpResponseBadRequest("Invalid content type.")

    documents = Document.objects.filter(_org_qs_filter(request)).order_by("-created_at")[:100]
    return render(request, "core/partials/doc_link_picker.html", {
        "documents": documents,
        "content_type_label": content_type_label,
        "object_id": object_id,
    })


@login_required
@require_POST
def create_doc_link(request):
    """Create a DocLink between a document and a target entity. Returns updated linked-documents section."""
    document_id = request.POST.get("document_id")
    content_type_label = request.POST.get("content_type_label")
    object_id = request.POST.get("object_id")

    if not all([document_id, content_type_label, object_id]):
        return HttpResponseBadRequest("Missing required fields.")

    if content_type_label not in _LINKABLE_MODELS:
        return HttpResponseBadRequest("Invalid content type.")

    app_label, model_name = _LINKABLE_MODELS[content_type_label]
    try:
        ct = ContentType.objects.get(app_label=app_label, model=model_name)
    except ContentType.DoesNotExist:
        return HttpResponseBadRequest("Content type not found.")

    document = get_object_or_404(Document, pk=document_id)

    DocLink.objects.get_or_create(
        document=document,
        content_type=ct,
        object_id=object_id,
        defaults={"created_by": request.user if request.user.is_authenticated else None},
    )

    if content_type_label == "prospect":
        prospect = get_object_or_404(Prospect, pk=object_id)
        doc_links = DocLink.objects.filter(
            content_type=ct,
            object_id=object_id,
        ).select_related("document", "created_by").order_by("-created_at")
        return render(request, "core/partials/linked_documents.html", {
            "entity": prospect,
            "doc_links": doc_links,
            "content_type_label": content_type_label,
        })

    return HttpResponse(status=204)


@login_required
@require_POST
def delete_doc_link(request, pk):
    """Delete a DocLink record and re-render the linked documents section."""
    link = get_object_or_404(DocLink, pk=pk)
    ct = link.content_type
    object_id = link.object_id
    content_type_label = ct.model

    link.delete()

    if content_type_label == "prospect":
        prospect = get_object_or_404(Prospect, pk=object_id)
        doc_links = DocLink.objects.filter(
            content_type=ct,
            object_id=object_id,
        ).select_related("document", "created_by").order_by("-created_at")
        return render(request, "core/partials/linked_documents.html", {
            "entity": prospect,
            "doc_links": doc_links,
            "content_type_label": content_type_label,
        })

    return HttpResponse(status=204)


@login_required
@require_GET
def drillholes(request):
    Drillhole = _get_model("core", "Drillhole")
    if Drillhole:
        qs = Drillhole.objects.filter(_org_qs_filter(request)).order_by("-created_at")
        page = _paginate(qs, request)
    else:
        qs, page = [], None
    return render(
        request,
        "core/drillholes.html",
        {"page": page, "model_exists": Drillhole is not None},
    )


@login_required
@require_GET
def tenements(request):
    Tenement = _get_model("core", "Tenement")
    if Tenement:
        qs = Tenement.objects.filter(_org_qs_filter(request)).order_by("-created_at")
        page = _paginate(qs, request)
    else:
        qs, page = [], None
    return render(
        request,
        "core/tenements.html",
        {"page": page, "model_exists": Tenement is not None},
    )


# ---------- AI / Map / Utilities ----------


@login_required
@require_GET
def ai_insights(request):
    """
    Placeholder page for AI features (report generation, summarization, etc.).
    """
    # You can pass recent docs/projects for prompts, etc.
    org_filter = _org_qs_filter(request)
    return render(
        request,
        "core/ai_insights.html",
        {
            "recent_docs": Document.objects.filter(org_filter).order_by("-created_at")[:12],
            "recent_projects": Process.objects.filter(org_filter).order_by("-created_at")[:8],
            "recent_reports": SavedReport.objects.filter(org_filter).select_related("process").order_by("-created_at")[:10],
        },
    )


@login_required
@require_GET
def map_view(request):
    """
    Simple map page that we should consider wiring to Leaflet / PostGIS endpoints.
    """
    return render(request, "core/map.html")


@require_GET
def healthcheck(request):
    """
    Lightweight container health endpoint (used by k8s/docker healthchecks later).
    """
    return JsonResponse({"status": "ok"})




@login_required
@require_GET
def project_report_pdf(request, process_id: str):
    org_filter = _org_qs_filter(request)
    if not Process.objects.filter(org_filter, pk=process_id).exists():
        raise Http404("Project not found")

    clearance_level = _get_clearance_level(request)
    md_text = _get_cached_report_md(process_id, clearance_level)
    process = Process.objects.get(pk=process_id)

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)

    styles = getSampleStyleSheet()
    # Custom styles
    h1 = ParagraphStyle('h1', parent=styles['Heading1'], textColor=colors.HexColor('#0e7490'), spaceAfter=10)
    h2 = ParagraphStyle('h2', parent=styles['Heading2'], textColor=colors.HexColor('#155e75'), spaceAfter=6)
    h3 = ParagraphStyle('h3', parent=styles['Heading3'], textColor=colors.HexColor('#1e4d5c'), spaceAfter=4)
    body = ParagraphStyle('body', parent=styles['Normal'], spaceAfter=6, leading=16)
    bullet = ParagraphStyle('bullet', parent=styles['Normal'], leftIndent=20, spaceAfter=4,
                             bulletIndent=10, leading=16)

    story = []
    for line in md_text.splitlines():
        if line.startswith('### '):
            story.append(Paragraph(line[4:], h3))
        elif line.startswith('## '):
            story.append(Paragraph(line[3:], h2))
        elif line.startswith('# '):
            story.append(Paragraph(line[2:], h1))
        elif line.startswith('- ') or line.startswith('* '):
            story.append(Paragraph(f'• {line[2:]}', bullet))
        elif re.match(r'^\d+\. ', line):
            story.append(Paragraph(re.sub(r'^\d+\. ', '', line), bullet))
        elif line.strip() == '':
            story.append(Spacer(1, 8))
        else:
            story.append(Paragraph(line, body))

    doc.build(story)
    buf.seek(0)
    slug = re.sub(r'[^\w-]', '_', process.name or str(process_id))
    response = HttpResponse(buf.getvalue(), content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{slug}_report.pdf"'
    return response

@login_required
@require_GET
def project_report_docx(request, process_id: str):
    org_filter = _org_qs_filter(request)
    if not Process.objects.filter(org_filter, pk=process_id).exists():
        raise Http404("Project not found")

    clearance_level = _get_clearance_level(request)
    md_text = _get_cached_report_md(process_id, clearance_level)
    process = Process.objects.get(pk=process_id)

    doc = DocxDocument()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for line in md_text.splitlines():
        if line.startswith("### "):
            p = doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            p = doc.add_heading(line[2:], level=1)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif re.match(r"^\d+\. ", line):
            doc.add_paragraph(re.sub(r"^\d+\. ", "", line), style="List Number")
        elif line.strip() == "":
            doc.add_paragraph("")
        else:
            doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    slug = re.sub(r"[^\w-]", "_", process.name or str(process_id))
    response = HttpResponse(
        buf.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    response["Content-Disposition"] = f'attachment; filename="{slug}_report.docx"'
    return response

@login_required
def document_analysis_detail(request, pk):
    document = get_object_or_404(Document, pk=pk)
    if not request.user.is_superuser:
        if (
            hasattr(request.user, 'profile')
            and request.user.profile.organisation
            and document.organisation
            and document.organisation != request.user.profile.organisation
        ):
            raise PermissionDenied

    analysis_text = getattr(document, "analysis_text", "") or "No insights available yet."

    return render(request, "core/document_analysis_detail.html", {
        "document": document,
        "analysis": analysis_text,
    })

@login_required
@require_POST
def save_document_analysis(request, pk):
    return HttpResponse("Save document analysis placeholder")


# ---------- GeoJSON API Endpoints for Map Viewer ----------


@login_required
@require_GET
def geojson_projects(request):
    """
    GeoJSON endpoint for Process (projects/operations) with spatial data
    Returns all processes with geometry for da map
    """
    from django.core.serializers import serialize
    from .models import Process

    # Only include processes with geometry
    processes = Process.objects.filter(
        _org_qs_filter(request), geom__isnull=False
    ).select_related('organisation')

    if not processes.exists():
        return JsonResponse({"type": "FeatureCollection", "features": []})

    # Use GeoDjangos built in serialiser (lets us take coordinates and translate into GeoJSOn text for geodata)
    geojson_data = serialize(
        'geojson',
        processes,
        geometry_field='geom',
        fields=('name', 'mode', 'commodity', 'organisation')
    )

    # Parse and return as JSON (serialise returns a string )
    import json
    return JsonResponse(json.loads(geojson_data), safe=False)


@login_required
@require_GET
def geojson_tenements(request):
    """
    GeoJSON endpoint for Tenement boundaries.
    Returns all tenements with geometry for map visualisation
    """
    from django.core.serializers import serialize
    from .models import Tenement

    tenements = Tenement.objects.filter(
        _org_qs_filter(request), geom__isnull=False
    ).select_related('organisation', 'process')

    if not tenements.exists():
        return JsonResponse({"type": "FeatureCollection", "features": []})

    geojson_data = serialize(
        'geojson',
        tenements,
        geometry_field='geom',
        fields=('name', 'organisation', 'process')
    )

    import json
    return JsonResponse(json.loads(geojson_data), safe=False)


@login_required
@require_GET
def geojson_prospects(request):
    """
    GeoJSON endpoint for Prospect locations
    Returns all prospects with geometry 
    """
    from django.core.serializers import serialize
    from .models import Prospect

    prospects = Prospect.objects.filter(
        _org_qs_filter(request), geom__isnull=False
    ).select_related('organisation', 'process')

    if not prospects.exists():
        return JsonResponse({"type": "FeatureCollection", "features": []})

    geojson_data = serialize(
        'geojson',
        prospects,
        geometry_field='geom',
        fields=('name', 'organisation', 'process')
    )

    import json
    return JsonResponse(json.loads(geojson_data), safe=False)


@login_required
@require_GET
def geojson_drillholes(request):
    """
    GeoJSON endpoint for Drillhole collar locations
    Returns all drillholes with collar locations 
    """
    from django.core.serializers import serialize
    from .models import Drillhole

    drillholes = Drillhole.objects.filter(
        _org_qs_filter(request), collar_location__isnull=False
    ).select_related('organisation', 'process')

    if not drillholes.exists():
        return JsonResponse({"type": "FeatureCollection", "features": []})

    geojson_data = serialize(
        'geojson',
        drillholes,
        geometry_field='collar_location',
        fields=('name', 'depth', 'azimuth', 'dip', 'organisation', 'process')
    )

    import json
    return JsonResponse(json.loads(geojson_data), safe=False)

# ---------- AI Report Generation & Document Analysis Pages ----------

@login_required
def report_list_page(request):
    clearance_rank = {"PUBLIC": 0, "INTERNAL": 1, "CONFIDENTIAL": 2, "JORC_APPROVED": 3}
    user_clearance = _get_clearance_level(request)
    user_rank = clearance_rank.get(user_clearance, 0)

    accessible_levels = [lvl for lvl, rank in clearance_rank.items() if rank <= user_rank]
    org_filter = _org_qs_filter(request)
    recent_reports = (
        SavedReport.objects
        .filter(org_filter, clearance_level__in=accessible_levels)
        .select_related("process")
        .order_by("-created_at")[:20]
    )
    recent_projects = Process.objects.filter(org_filter).order_by("-created_at")[:20]
    all_documents = Document.objects.filter(org_filter).select_related("process").order_by("-created_at")

    return render(request, "core/report_list.html", {
        "recent_reports":  recent_reports,
        "recent_projects": recent_projects,
        "all_documents":   all_documents,
    })


@login_required
@role_required(
    UserProfile.RoleChoices.GEOLOGIST_EXPL,
    UserProfile.RoleChoices.FIELD_LEAD,
    UserProfile.RoleChoices.DATA_MANAGER,
    UserProfile.RoleChoices.GEOLOGIST_MINE,
    UserProfile.RoleChoices.METALLURGIST,
    UserProfile.RoleChoices.OPERATIONS_MANAGER,
    UserProfile.RoleChoices.ADMIN,
)
def generate_report(request):
    """
    POST: generate (or retrieve cached) report for a process and redirect to the editor.
    GET:  redirect back to the report list.
    """
    if request.method != "POST":
        return redirect("report_list")

    process_id   = request.POST.get("process_id", "").strip()
    report_title = request.POST.get("report_title", "").strip()

    if not process_id:
        messages.error(request, "No project selected.")
        return redirect("report_list")

    org_filter = _org_qs_filter(request)
    if not Process.objects.filter(org_filter, pk=process_id).exists():
        raise Http404("Project not found")

    clearance_level = _get_clearance_level(request)
    try:
        process = get_object_or_404(Process, org_filter, pk=process_id)
        md = _get_cached_report_md(process_id, clearance_level)
    except Exception as e:
        log.error("Report generation failed during generate_report: %s", e)
        messages.error(request, f"Report generation failed: {e}")
        return redirect("report_list")

    import hashlib
    title = report_title or f"{process.name or 'Project'} Report"
    existing = SavedReport.objects.filter(
        process=process, title=title
    ).order_by("-version_number").first()

    if existing:
        SavedReport.create_version(
            parent=existing,
            content_md=md,
            user=request.user,
            reason=SavedReport.ChangeReason.REGENERATED,
        )
    else:
        SavedReport.objects.create(
            process=process,
            organisation=process.organisation,
            title=title,
            content_md=md,
            content_hash=hashlib.sha256(md.encode()).hexdigest(),
            created_by=request.user,
            version_number=1,
            change_reason=SavedReport.ChangeReason.GENERATED,
        )

    return redirect(reverse("report_editor", kwargs={"process_id": process_id}))


@login_required
def report_editor(request, process_id):
    """
    Serve the report editor page for a process
    loads markdown from cache (generated by generate_report before it redirects into here)
    """
    org_filter = _org_qs_filter(request)
    try:
        process = Process.objects.filter(org_filter).select_related("organisation").get(pk=process_id)
    except Process.DoesNotExist:
        raise Http404("Project not found")

    clearance_level = _get_clearance_level(request)
    try:
        md = _get_cached_report_md(str(process_id), clearance_level)
    except Exception as e:
        log.error("Report editor cache miss for process %s: %s", process_id, e)
        md = f"# {process.name or 'Project'} Report\n\nReport generation failed: {e}"

    custom_title  = request.GET.get("title", "").strip()
    default_title = custom_title or f"{process.name or 'Project'} Report"

    return render(request, "core/report_editor.html", {
        "process": process,
        "markdown_content": md,
        "default_title": default_title,
        "saved_report": None,
        "save_url": reverse("save_report"),
        "export_url": reverse("export_report"),
    })


@login_required
def saved_report_editor(request, report_id):
    """ serve the report editor page for an existing saved report """
    report = get_object_or_404(
        SavedReport.objects.select_related("process", "organisation"),
        pk=report_id,
    )

    if not request.user.is_superuser:
        if (
            hasattr(request.user, 'profile')
            and request.user.profile.organisation
            and report.organisation
            and report.organisation != request.user.profile.organisation
        ):
            raise PermissionDenied

    user_clearance = _get_clearance_level(request)
    clearance_rank = {"PUBLIC": 0, "INTERNAL": 1, "CONFIDENTIAL": 2, "JORC_APPROVED": 3}
    if clearance_rank.get(user_clearance, 0) < clearance_rank.get(report.clearance_level, 1):
        raise PermissionDenied

    return render(request, "core/report_editor.html", {
        "process": report.process,
        "markdown_content": report.content_md,
        "default_title": report.title,
        "saved_report": report,
        "save_url": reverse("update_saved_report", kwargs={"report_id": report_id}),
        "export_url": reverse("export_report"),
    })


@login_required
@require_POST
def save_report(request):
    """
    create a new SavedReport record from the user/editors current content
    returns JSON: {success: true, report_id: "...", redirect_url: "..."}
    """
    process_id = request.POST.get("process_id", "").strip()
    title      = request.POST.get("title", "").strip()
    content_md = request.POST.get("content_md", "").strip()

    if not title:
        return JsonResponse({"success": False, "error": "Title is required."}, status=400)
    if not content_md:
        return JsonResponse({"success": False, "error": "Report content is empty."}, status=400)

    process = None
    organisation = None
    if process_id:
        try:
            process = Process.objects.select_related("organisation").get(pk=process_id)
            organisation = process.organisation
        except Process.DoesNotExist:
            pass

    clearance_level = _get_clearance_level(request)
    created_by = request.user if request.user.is_authenticated else None

    existing = SavedReport.objects.filter(
        process=process, title=title
    ).order_by("-version_number").first()

    if existing:
        report = SavedReport.create_version(
            parent=existing,
            content_md=content_md,
            user=created_by,
            reason=SavedReport.ChangeReason.MANUAL_EDIT,
        )
    else:
        import hashlib
        report = SavedReport.objects.create(
            process=process,
            organisation=organisation,
            title=title,
            content_md=content_md,
            content_hash=hashlib.sha256(content_md.encode()).hexdigest(),
            clearance_level=clearance_level,
            created_by=created_by,
            version_number=1,
            change_reason=SavedReport.ChangeReason.GENERATED,
        )
    return JsonResponse({
        "success": True,
        "report_id": str(report.id),
        "version_number": report.version_number,
        "redirect_url": reverse("saved_report_editor", kwargs={"report_id": report.id}),
    })


@login_required
@require_POST
def update_saved_report(request, report_id):
    """
    overwrite of an existing SavedReport title and content
    Returns JSON: {success: true}
    """
    report = get_object_or_404(SavedReport, pk=report_id)

    is_admin = hasattr(request.user, "profile") and request.user.profile.role == "ADMIN"
    if report.created_by != request.user and not is_admin:
        return JsonResponse({"success": False, "error": "Permission denied."}, status=403)

    title      = request.POST.get("title", "").strip()
    content_md = request.POST.get("content_md", "").strip()

    if not title:
        return JsonResponse({"success": False, "error": "Title is required."}, status=400)
    if not content_md:
        return JsonResponse({"success": False, "error": "Report content is empty."}, status=400)

    new_version = SavedReport.create_version(
        parent=report,
        content_md=content_md,
        user=request.user,
        reason=SavedReport.ChangeReason.MANUAL_EDIT,
    )
    return JsonResponse({
        "success": True,
        "new_version_id": str(new_version.id),
        "version_number": new_version.version_number,
    })

@login_required
def report_history(request, process_id):
    """Show all versions of reports for a process, grouped by title."""
    # Get the latest version of each distinct report title
    org_filter = _org_qs_filter(request)
    reports = (
        SavedReport.objects
        .filter(org_filter, process_id=process_id)
        .order_by("title", "-version_number")
    )
    # Group by title to show each report with its version chain
    from itertools import groupby
    grouped = {
        title: list(versions)
        for title, versions in groupby(reports, key=lambda r: r.title)
    }
    return render(request, "core/report_history.html", {"grouped": grouped, "process_id": process_id})

@login_required
def report_version_detail(request, report_id):
    """View a specific report version."""
    report = get_object_or_404(SavedReport, pk=report_id)
    all_versions = SavedReport.objects.filter(
        process=report.process, title=report.title
    ).order_by("-version_number")

    log_audit(
        user=request.user,
        action=AuditLog.ActionType.VIEW,
        obj=report,
        description=f"Viewed '{report.title}' v{report.version_number}",
        ip_address=request.META.get("REMOTE_ADDR"),
        user_agent=request.META.get("HTTP_USER_AGENT", ""),
    )
    return render(request, "core/report_version_detail.html", {
        "report": report,
        "all_versions": all_versions,
    })

@login_required
@require_GET
def all_reports_history(request):
    """Show all saved reports grouped by project."""
    from itertools import groupby

    org_filter = _org_qs_filter(request)
    reports = (
        SavedReport.objects
        .filter(org_filter)
        .select_related("process")
        .order_by("process__name", "title", "-version_number")
    )

    grouped = {}
    for report in reports:
        project_name = report.process.name if report.process else "No Project"
        process_id = str(report.process.id) if report.process else None
        if project_name not in grouped:
            grouped[project_name] = {"process_id": process_id, "titles": {}}
        if report.title not in grouped[project_name]["titles"]:
             grouped[project_name]["titles"][report.title] = []
        grouped[project_name]["titles"][report.title].append(report)

    return render(request, "core/all_reports_history.html", {"grouped": grouped})


@login_required
@require_POST
def export_report(request):
    """
    Export the  current markdown content as PDF or DOCX.
    POST params:
        format — "pdf" or "docx"
        content_md — the markdown string to render
        title — used as the filename 
    """
    fmt       = request.POST.get("format", "pdf").lower()
    md_text   = request.POST.get("content_md", "")
    title     = request.POST.get("title", "report")
    slug      = re.sub(r"[^\w-]", "_", title)

    if fmt == "pdf":
        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4,
                                leftMargin=2*cm, rightMargin=2*cm,
                                topMargin=2*cm, bottomMargin=2*cm)
        styles = getSampleStyleSheet()
        h1     = ParagraphStyle("h1", parent=styles["Heading1"], textColor=colors.HexColor("#0e7490"), spaceAfter=10)
        h2     = ParagraphStyle("h2", parent=styles["Heading2"], textColor=colors.HexColor("#155e75"), spaceAfter=6)
        h3     = ParagraphStyle("h3", parent=styles["Heading3"], textColor=colors.HexColor("#1e4d5c"), spaceAfter=4)
        body   = ParagraphStyle("body", parent=styles["Normal"], spaceAfter=6, leading=16)
        bullet = ParagraphStyle("bullet", parent=styles["Normal"], leftIndent=20, spaceAfter=4, bulletIndent=10, leading=16)

        story = []
        for line in md_text.splitlines():
            if line.startswith("### "):
                story.append(Paragraph(line[4:], h3))
            elif line.startswith("## "):
                story.append(Paragraph(line[3:], h2))
            elif line.startswith("# "):
                story.append(Paragraph(line[2:], h1))
            elif line.startswith("- ") or line.startswith("* "):
                story.append(Paragraph(f"• {line[2:]}", bullet))
            elif re.match(r"^\d+\. ", line):
                story.append(Paragraph(re.sub(r"^\d+\. ", "", line), bullet))
            elif line.strip() == "":
                story.append(Spacer(1, 8))
            else:
                story.append(Paragraph(line, body))

        doc.build(story)
        buf.seek(0)
        response = HttpResponse(buf.getvalue(), content_type="application/pdf")
        response["Content-Disposition"] = f'attachment; filename="{slug}_report.pdf"'
        return response

    if fmt == "docx":
        doc = DocxDocument()
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        for line in md_text.splitlines():
            if line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:], style="List Bullet")
            elif re.match(r"^\d+\. ", line):
                doc.add_paragraph(re.sub(r"^\d+\. ", "", line), style="List Number")
            elif line.strip() == "":
                doc.add_paragraph("")
            else:
                doc.add_paragraph(line)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        response = HttpResponse(
            buf.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        response["Content-Disposition"] = f'attachment; filename="{slug}_report.docx"'
        return response

    return JsonResponse({"error": "Invalid format. Use 'pdf' or 'docx'."}, status=400)


@login_required
def report_detail(request, report_id):
    return render(request, "core/report_detail.html", {
        "report_id": report_id,
    })


@login_required
def document_analysis_page(request):
    return render(request, "core/document_analysis.html", {
        "recent_docs": Document.objects.filter(_org_qs_filter(request)).select_related("process").order_by("-created_at"),
    })


@login_required
def analyze_document(request, pk):
    document = get_object_or_404(Document, pk=pk)

    if not request.user.is_superuser:
        if (
            hasattr(request.user, 'profile')
            and request.user.profile.organisation
            and document.organisation
            and document.organisation != request.user.profile.organisation
        ):
            raise PermissionDenied

    text = (document.extracted_text or "").strip()
    if not text:
        messages.error(request, "This document has no extracted text to analyse.")
        return redirect("document_analysis_page")

    try:
        client = GraniteClient()

        prompt = f"""
You are analysing a mining/exploration document.

Provide:
- A short summary
- Key insights
- Risks or issues
- Important findings
- Suggested next steps

Return the analysis in this exact format:

## Summary
...

## Key Insights
- ...

## Risks
- ...

## Recommended Actions
- ...
Document title: {document.title}

Document text:
{text[:12000]}
"""

        analysis_text = client.complete(prompt)

        document.analysis_text = analysis_text
        document.save(update_fields=["analysis_text"])

        messages.success(request, "Analysis complete.")
        return redirect("document_analysis_detail", pk=document.pk)

    except Exception as e:
        messages.error(request, f"Analysis failed: {e}")
        return redirect("document_analysis_page")

@login_required
@require_GET
def export_document_analysis(request, pk):
    document = get_object_or_404(Document, pk=pk)

    md_text = (document.analysis_text or "").strip()
    if not md_text:
        return JsonResponse({"error": "No analysis available to export."}, status=400)

    fmt = request.GET.get("format", "pdf").lower()
    title = f"{document.title} Analysis"
    slug = re.sub(r"[^\w-]", "_", title)

    if fmt == "pdf":
        buf = io.BytesIO()
        doc = SimpleDocTemplate(
            buf,
            pagesize=A4,
            leftMargin=2*cm,
            rightMargin=2*cm,
            topMargin=2*cm,
            bottomMargin=2*cm,
        )
        styles = getSampleStyleSheet()
        h1 = ParagraphStyle("h1", parent=styles["Heading1"], textColor=colors.HexColor("#0e7490"), spaceAfter=10)
        h2 = ParagraphStyle("h2", parent=styles["Heading2"], textColor=colors.HexColor("#155e75"), spaceAfter=6)
        h3 = ParagraphStyle("h3", parent=styles["Heading3"], textColor=colors.HexColor("#1e4d5c"), spaceAfter=4)
        body = ParagraphStyle("body", parent=styles["Normal"], spaceAfter=6, leading=16)
        bullet = ParagraphStyle("bullet", parent=styles["Normal"], leftIndent=20, spaceAfter=4, bulletIndent=10, leading=16)

        story = []
        for line in md_text.splitlines():
            if line.startswith("### "):
                story.append(Paragraph(line[4:], h3))
            elif line.startswith("## "):
                story.append(Paragraph(line[3:], h2))
            elif line.startswith("# "):
                story.append(Paragraph(line[2:], h1))
            elif line.startswith("- ") or line.startswith("* "):
                story.append(Paragraph(f"• {line[2:]}", bullet))
            elif re.match(r"^\d+\. ", line):
                story.append(Paragraph(re.sub(r"^\d+\. ", "", line), bullet))
            elif line.strip() == "":
                story.append(Spacer(1, 8))
            else:
                story.append(Paragraph(line, body))

        doc.build(story)
        buf.seek(0)
        response = HttpResponse(buf.getvalue(), content_type="application/pdf")
        response["Content-Disposition"] = f'attachment; filename="{slug}_analysis.pdf"'
        return response

    if fmt == "docx":
        doc = DocxDocument()
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

        for line in md_text.splitlines():
            if line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:], style="List Bullet")
            elif re.match(r"^\d+\. ", line):
                doc.add_paragraph(re.sub(r"^\d+\. ", "", line), style="List Number")
            elif line.strip() == "":
                doc.add_paragraph("")
            else:
                doc.add_paragraph(line)

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        response = HttpResponse(
            buf.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        response["Content-Disposition"] = f'attachment; filename="{slug}_analysis.docx"'
        return response

    return JsonResponse({"error": "Invalid format. Use 'pdf' or 'docx'."}, status=400)